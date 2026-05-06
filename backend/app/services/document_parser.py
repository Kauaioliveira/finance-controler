from __future__ import annotations

from io import BytesIO
from pathlib import Path

from docx import Document as DocxDocument
from pypdf import PdfReader

from app.core.exceptions import DocumentProcessingError, ValidationError


SUPPORTED_DOCUMENT_EXTENSIONS = {
    ".pdf": "application/pdf",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ".txt": "text/plain",
    ".md": "text/markdown",
}


class DocumentParser:
    def parse(self, filename: str, payload: bytes) -> tuple[str, str]:
        extension = Path(filename).suffix.lower()
        if extension not in SUPPORTED_DOCUMENT_EXTENSIONS:
            raise ValidationError(
                "Formato de arquivo nao suportado. Use PDF, DOCX, TXT ou MD."
            )
        if not payload:
            raise ValidationError("O arquivo enviado esta vazio.")

        parser = {
            ".pdf": self._parse_pdf,
            ".docx": self._parse_docx,
            ".txt": self._parse_text,
            ".md": self._parse_text,
        }[extension]

        try:
            extracted_text = parser(payload)
        except DocumentProcessingError:
            raise
        except Exception as exc:
            raise DocumentProcessingError(
                f"Nao foi possivel processar o arquivo {filename}."
            ) from exc

        normalized_text = extracted_text.strip()
        if not normalized_text:
            raise DocumentProcessingError(
                f"Nao foi possivel extrair texto util do arquivo {filename}."
            )

        return SUPPORTED_DOCUMENT_EXTENSIONS[extension], normalized_text

    def _parse_pdf(self, payload: bytes) -> str:
        reader = PdfReader(BytesIO(payload))
        parts = []
        for page in reader.pages:
            parts.append(page.extract_text() or "")
        return "\n\n".join(part for part in parts if part.strip())

    def _parse_docx(self, payload: bytes) -> str:
        document = DocxDocument(BytesIO(payload))
        parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
        for table in document.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    parts.append(row_text)
        return "\n".join(parts)

    def _parse_text(self, payload: bytes) -> str:
        for encoding in ("utf-8", "utf-8-sig", "latin-1"):
            try:
                return payload.decode(encoding)
            except UnicodeDecodeError:
                continue
        raise DocumentProcessingError("Nao foi possivel decodificar o arquivo texto.")


document_parser = DocumentParser()
